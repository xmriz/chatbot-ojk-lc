import fitz
from docx import Document as DocxDocument
from openpyxl import load_workbook
from langchain_core.documents import Document
import pandas as pd
from paddleocr import PaddleOCR
import os

# ==================== EXTRACT TEXT FROM PDF ====================


def extract_text_and_images_from_page(doc, page, ocr, treshold):
    text = page.get_text()
    image_text = ""
    image_list = page.get_images(full=True)
    # Iterate through all images found on the page
    for image_info in image_list:
        xref = image_info[0]
        image_dict = doc.extract_image(xref)
        image_bytes = image_dict['image']

        if image_bytes is not None:
            # Use PaddleOCR to extract text from the image
            ocr_result = ocr.ocr(image_bytes)
            # Check if OCR result is valid before processing
            if ocr_result and ocr_result != [None]:
                for result in ocr_result:
                    for res in result:
                        text_tuple = res[1]
                        text_string = text_tuple[0]
                        # For confidence threshold
                        text_confidence = text_tuple[1]
                        if text_confidence > treshold:
                            image_text += text_string + '\n'
    # Combine page text and image text
    return text + "\n" + image_text


def extract_text_from_pdf(file_path, ocr, treshold):
    # Load the PDF file
    doc = fitz.open(file_path)
    text = ""
    # Iterate through all pages in the PDF
    for page_num in range(len(doc)):
        page = doc[page_num]
        # Extract text and images from the page
        page_text = extract_text_and_images_from_page(doc, page, ocr, treshold)
        text += page_text + "\n"
    return text


# ==================== EXTRACT TEXT FROM DOCX ====================

def extract_text_from_docx(file_path):
    doc = DocxDocument(file_path)
    text = ''
    for para in doc.paragraphs:
        text += para.text + '\n'
    return text


# ==================== EXTRACT TEXT FROM EXCEL ====================

def extract_text_from_excel(excel_path):
    wb = load_workbook(excel_path)
    text_data = []

    for sheet in wb.worksheets:
        for row in sheet.iter_rows(values_only=True):
            for cell in row:
                if cell is not None:
                    text_data.append(str(cell))

    return ' '.join(text_data)


# # ==================== MAIN ====================


def extract_all_documents_in_directory(documents_dir, metadata_path, treshold=0.98):
    ocr = PaddleOCR(use_angle_cls=True, lang='id', show_log=False)

    docs = []

    df_metadata = pd.read_csv(metadata_path)

    for file in os.listdir(documents_dir):
        file_path = os.path.join(documents_dir, file)
        file_metadata = df_metadata[df_metadata['new_filename'] == file]
        if file.endswith('.pdf'):
            text = extract_text_from_pdf(file_path, ocr, treshold)
            document = Document(
                page_content=text,
                metadata={
                    "file_name": file_metadata['new_filename'].values[0],
                    "title": file_metadata['title'].values[0],
                    "sector": file_metadata['sektor'].values[0],
                    "subsector": file_metadata['subsektor'].values[0],
                    "regulation_type": file_metadata['jenis_regulasi'].values[0],
                    "regulation_number": file_metadata['nomor_regulasi'].values[0],
                    "effective_date": file_metadata['tanggal_berlaku'].values[0],
                    "file_url": file_metadata['file_url'].values[0],
                }
            )
            docs.append(document)
        elif file.endswith('.xlsm') or file.endswith('.xlsx') or file.endswith('.xls'):
            text = extract_text_from_excel(file_path)
            document = Document(
                page_content=text,
                metadata={
                    "file_name": file_metadata['new_filename'].values[0],
                    "title": file_metadata['title'].values[0],
                    "sector": file_metadata['sektor'].values[0],
                    "subsector": file_metadata['subsektor'].values[0],
                    "regulation_type": file_metadata['jenis_regulasi'].values[0],
                    "regulation_number": file_metadata['nomor_regulasi'].values[0],
                    "effective_date": file_metadata['tanggal_berlaku'].values[0],
                    "file_url": file_metadata['file_url'].values[0],
                }
            )
            docs.append(document)
        elif file.endswith('.docx'):
            text = extract_text_from_docx(file_path)
            document = Document(
                page_content=text,
                metadata={
                    "file_name": file_metadata['new_filename'].values[0],
                    "title": file_metadata['title'].values[0],
                    "sector": file_metadata['sektor'].values[0],
                    "subsector": file_metadata['subsektor'].values[0],
                    "regulation_type": file_metadata['jenis_regulasi'].values[0],
                    "regulation_number": file_metadata['nomor_regulasi'].values[0],
                    "effective_date": file_metadata['tanggal_berlaku'].values[0],
                    "file_url": file_metadata['file_url'].values[0],
                }
            )
            docs.append(document)
    print(f"Read {len(docs)} documents")
    return docs